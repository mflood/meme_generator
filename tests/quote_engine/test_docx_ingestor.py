from unittest.mock import MagicMock

import pytest

from quote_engine.docx_ingestor import DocxIngestor
from quote_engine.models import QuoteModel


@pytest.fixture
def mock_docx_file(tmp_path):
    from docx import Document

    doc = Document()
    doc.add_paragraph('"Life is beautiful" - Anonymous')
    doc.add_paragraph('"To be or not to be" - Shakespeare')

    path = tmp_path / "test.docx"
    doc.save(path)

    return path


def test_can_ingest():
    assert DocxIngestor.can_ingest("test.docx") == True
    assert DocxIngestor.can_ingest("test.pdf") == False
    assert DocxIngestor.can_ingest("test.txt") == False
    assert DocxIngestor.can_ingest("test.doc") == False


def test_parse(mock_docx_file):
    quotes = DocxIngestor.parse(mock_docx_file)
    assert len(quotes) == 2
    assert quotes[0].body == "Life is beautiful"
    assert quotes[0].author == "Anonymous"
    assert quotes[1].body == "To be or not to be"
    assert quotes[1].author == "Shakespeare"


def test_parse_invalid_lines(monkeypatch):
    from docx import Document

    mock_document = MagicMock()
    mock_document.paragraphs = [
        MagicMock(text="Invalid quote line"),
        MagicMock(text='"Valid quote" - Author'),
    ]

    def mock_document_constructor(*args, **kwargs):
        return mock_document

    monkeypatch.setattr(
        "quote_engine.docx_ingestor.Document", mock_document_constructor
    )

    mock_path = "fake_path.docx"
    quotes = DocxIngestor.parse(mock_path)
    assert len(quotes) == 1
    assert quotes[0].body == "Valid quote"
    assert quotes[0].author == "Author"


if __name__ == "__main__":
    pytest.main()
